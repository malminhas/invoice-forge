"""
Invoice Generator API Service

This FastAPI service provides RESTful endpoints to generate professional invoices in DOCX and PDF formats.
It leverages the core invoice generation logic from the invoice_generator.py script, exposed as an API.

## API Endpoints
- GET /: API root with basic information
- GET /version: API version and configuration information
- GET /example-client: Returns an HTML test client
- POST /generate-invoice: Generate and download invoice file
    
## Authentication
- No authentication is required for this version
- For production use, consider adding API keys or other authentication methods

## CORS
- CORS is enabled for all origins by default (suitable for development)
- For production, restrict origins to specific domains
- The API allows common HTTP methods and headers
- Credentials are supported for authenticated requests

## Installation
- pip install fastapi uvicorn pydantic python-docx pyyaml docopt

## Usage
- Start the server using uvicorn:
  ```
  # Basic usage (default: http://localhost:8000)
  uvicorn invoice_generator_api:app --reload
  
  # With custom host and port
  uvicorn invoice_generator_api:app --host 0.0.0.0 --port 8000 --reload
  
  # Production mode (disable reload)
  uvicorn invoice_generator_api:app --host 0.0.0.0 --port 8083
  
  # With verbose logging
  VERBOSE=True uvicorn invoice_generator_api:app --reload
  
  # Using Python module syntax
  python -m uvicorn invoice_generator_api:app --reload
  ```

- Access Swagger UI: http://localhost:8083/docs
- Access ReDoc UI: http://localhost:8083/redoc
- Send POST requests to /generate-invoice with appropriate JSON data

## Invoice Generation Process
1. The client sends invoice details as JSON to the /generate-invoice endpoint
2. The service validates the input data using Pydantic models
3. Invoice generation logic creates a DOCX document using python-docx
4. If PDF format is requested, the service converts DOCX to PDF using unoconv
5. The generated file is returned as a downloadable response
    
## Requirements
- FastAPI: pip install fastapi
- Uvicorn: pip install uvicorn
- Pydantic: pip install pydantic
- Python-DOCX: pip install python-docx
- PyYAML: pip install pyyaml
- Docopt: pip install docopt
    
For PDF conversion:
- LibreOffice: brew install libreoffice
- unoconv: brew install unoconv

## Example API Request (using curl)

curl -X POST "http://localhost:8083/generate-invoice?format=pdf" \\
    -H "Content-Type: application/json" \\
    -d '{
        "client_name": "Mike Smith",
        "client_address": "17 Poland St.\\nLondon\\nW1 1ZZ\\nU.K.",
        "services": ["AI Consultancy 29.03.25 (1 hour)", "Notes write up 29.03.25 (1 hour)"],
        "payment_terms_days": 30,
        "invoice_number": 1008,
        "invoice_date": "21.04.25",
        "company_name": "Fizzbuzz Consulting Ltd",
        "hourly_rate": 300,
        "vat_rate": 20,
        "account_number": "12345678",
        "sort_code": "12-34-56",
        "bank_address": "123 Bank St, London, UK",
        "company_number": "12345678",
        "vat_number": "GB123456789",
        "registered_address": "123 Business St, London, UK",
        "email": "contact@myconsulting.com",
        "contact_number": "07700 900123",
        "column_widths": [2.5, 3.5],
        "font_name": "DejaVu Sans",
        "icon_name": "FizzbuzzConsultingIcon.png"
    }' \\
    -o invoice.pdf
"""

import os
import tempfile
import logging
import yaml # type: ignore
import sys
from contextlib import asynccontextmanager
from typing import Dict, List, Union, Optional, Set # type: ignore
from fastapi import FastAPI, HTTPException, Query, status # type: ignore
from fastapi.responses import FileResponse # type: ignore
from fastapi.middleware.cors import CORSMiddleware # type: ignore
from pydantic import BaseModel, Field, field_validator # type: ignore
from pathlib import Path
from enum import Enum
from fastapi.openapi.utils import get_openapi # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches

# Configuration variables
VERBOSE = os.environ.get('VERBOSE', 'False').lower() in ('true', '1', 't')
LOG_FORMAT = '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
LOG_DATE_FORMAT = '%Y-%m-%d %H:%M:%S'
LOG_LEVEL = logging.DEBUG if VERBOSE else logging.INFO

# CORS Configuration - Customize for your needs
CORS_ORIGINS = os.environ.get('CORS_ORIGINS', '*').split(',')
CORS_ALLOW_CREDENTIALS = os.environ.get('CORS_ALLOW_CREDENTIALS', 'True').lower() in ('true', '1', 't')
CORS_ALLOW_METHODS = ["GET", "POST", "OPTIONS"]
CORS_ALLOW_HEADERS = ["*"]

# Setup logging configuration
logging.basicConfig(
    level=LOG_LEVEL,
    format=LOG_FORMAT,
    datefmt=LOG_DATE_FORMAT,
    stream=sys.stdout
)

# Get logger for this module
logger = logging.getLogger(__name__)

# Import functions from invoice_generator.py
from invoice_generator import (
    setup_logging,
    ensure_invoices_directory,
    get_output_path,
    convert_to_pdf,
    get_pdf_path,
    Document,
    WD_ALIGN_PARAGRAPH,
    Inches,
    Pt,
    OxmlElement,
    qn
)

# Define output format enum for better Swagger docs
class OutputFormat(str, Enum):
    DOCX = "docx"
    PDF = "pdf"

# Configure the invoice_generator module's logging
setup_logging(verbose=VERBOSE)

logger.info(f"Starting Invoice Generator API with VERBOSE={VERBOSE}")

# Lifespan event handler for startup and shutdown
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    logger.info("Invoice Generator API is starting up")
    logger.info(f"Verbose logging is {'enabled' if VERBOSE else 'disabled'}")
    logger.info(f"API documentation available at http://localhost:8083/docs")
    yield
    # Shutdown
    logger.info("Invoice Generator API is shutting down")

# Create FastAPI app with enhanced metadata
app = FastAPI(
    title="Invoice Generator API",
    description="""
This API allows you to generate professional invoices in DOCX and PDF formats.

## Features

* Generate invoices with customizable details
* Support for both DOCX and PDF output formats
* Automatic calculation of subtotals, VAT, and total amounts
* Professional formatting with company branding

## Usage

Send a POST request to /generate-invoice with your invoice details as JSON.

See the Schema section below for the required and optional fields.
    """,
    version="1.0.0",
    contact={
        "name": "Invoice Generator Support",
        "email": "support@example.com",
    },
    license_info={
        "name": "MIT License",
    },
    lifespan=lifespan,
    root_path=os.environ.get('ROOT_PATH', '')
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=CORS_ALLOW_CREDENTIALS,
    allow_methods=CORS_ALLOW_METHODS,
    allow_headers=CORS_ALLOW_HEADERS,
)

logger.info(f"CORS middleware configured with origins: {CORS_ORIGINS}")

# Ensure the invoices directory exists
invoices_dir = ensure_invoices_directory()
logger.info(f"Using invoices directory: {invoices_dir}")

class InvoiceDetails(BaseModel):
    """
    Invoice details model containing all fields required to generate an invoice
    """
    client_name: str = Field(
        ..., 
        description="Client's full name or company name",
        example="Mike Smith"
    )
    client_address: str = Field(
        ..., 
        description="Client's full address with line breaks as needed",
        example="17 Poland St.\nLondon\nW2 4ZZ\nU.K."
    )
    services: List[str] = Field(
        ..., 
        description="List of services with date (DD.MM.YY) and hours in parentheses, e.g. 'Service name 21.04.25 (2 hours)'",
        example=["AI Consultancy 29.03.25 (1 hour)", "Notes write up 29.03.25 (1 hour)"]
    )
    service_date: Optional[str] = Field(
        None,
        description="Date when the service was provided in DD.MM.YY format",
        example="21.04.25"
    )
    service_description: Optional[str] = Field(
        None,
        description="General description of the service provided",
        example="AI Consultancy and Documentation"
    )
    payment_terms_days: int = Field(
        ..., 
        description="Payment terms in days",
        example=30,
        gt=0
    )
    invoice_number: int = Field(
        ..., 
        description="Unique invoice number",
        example=1008,
        gt=0
    )
    invoice_date: Optional[str] = Field(
        None, 
        description="Invoice date in DD.MM.YY format. If not provided, today's date will be used",
        example="21.04.25"
    )
    company_name: str = Field(
        ..., 
        description="Your company name",
        example="My Consulting Ltd"
    )
    hourly_rate: float = Field(
        ..., 
        description="Hourly rate in GBP",
        example=300.0,
        gt=0
    )
    vat_rate: float = Field(
        20.0, 
        description="VAT rate as a percentage",
        example=20.0,
        ge=0
    )
    account_number: str = Field(
        ..., 
        description="Bank account number",
        example="12345678"
    )
    sort_code: str = Field(
        ..., 
        description="Bank sort code",
        example="12-34-56"
    )
    bank_address: str = Field(
        ..., 
        description="Bank address",
        example="123 Bank St, London, UK"
    )
    company_number: str = Field(
        ..., 
        description="Company registration number",
        example="12345678"
    )
    vat_number: str = Field(
        ..., 
        description="VAT registration number",
        example="GB123456789"
    )
    registered_address: str = Field(
        ..., 
        description="Company registered address",
        example="123 Business St, London, UK"
    )
    email: str = Field(
        ..., 
        description="Contact email address",
        example="contact@myconsulting.com"
    )
    contact_number: str = Field(
        ..., 
        description="Contact phone number",
        example="07700 900123"
    )
    column_widths: List[float] = Field(
        [2.5, 3.5], 
        description="Document column widths in inches [left, right]",
        example=[2.5, 3.5]
    )
    font_name: str = Field(
        "DejaVu Sans", 
        description="Font name to use in the document",
        example="DejaVu Sans"
    )
    icon_name: str = Field(
        "DioramaConsultingIcon.png", 
        description="Company icon/logo filename (must be in the same directory)",
        example="DioramaConsultingIcon.png"
    )
    icon_data: Optional[str] = Field(
        None,
        description="Base64 encoded image data for the icon",
        example="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAA..."
    )
    paid: Optional[bool] = Field(
        False,
        description="Whether the invoice has been paid",
        example=False
    )
    
    @field_validator('services')
    @classmethod
    def validate_services(cls, v):
        """Validate that services contain date and hours in parentheses"""
        import re
        for service in v:
            # Check for hours in parentheses
            if not re.search(r'\(\d+\.?\d*\s*hours?\)', service):
                raise ValueError(f"Service '{service}' must include hours in parentheses, e.g. '(2 hours)'")
            
            # Check for date in format DD.MM.YY or DD.MM.YYYY
            if not re.search(r'\d{1,2}\.\d{1,2}\.\d{2,4}', service):
                raise ValueError(f"Service '{service}' must include a date in format DD.MM.YY, e.g. '21.04.25'")
        return v
        
    class Config:
        json_schema_extra = {
            "example": {
                "client_name": "Mike Smith",
                "client_address": "17 Poland St.\nLondon\nW2 4ZZ\nU.K.",
                "services": ["AI Consultancy 29.03.25 (1 hour)", "Notes write up 29.03.25 (1 hour)"],
                "service_date": "29.03.25",
                "service_description": "AI Consultancy and Documentation Services",
                "payment_terms_days": 30,
                "invoice_number": 1008,
                "invoice_date": "21.04.25",
                "company_name": "My Consulting Ltd",
                "hourly_rate": 300,
                "vat_rate": 20,
                "account_number": "12345678",
                "sort_code": "12-34-56",
                "bank_address": "123 Bank St, London, UK",
                "company_number": "12345678",
                "vat_number": "GB123456789",
                "registered_address": "123 Business St, London, UK",
                "email": "contact@myconsulting.com",
                "contact_number": "07700 900123",
                "column_widths": [2.5, 3.5],
                "font_name": "DejaVu Sans",
                "icon_name": "DioramaConsultingIcon.png",
                "icon_data": "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+P+/HgAFdwI2hN4pSgAAAABJRU5ErkJggg=="
            }
        }

def generate_invoice_document(details: Dict, output_path: Path, generate_pdf: bool = False):
    """
    Generate an invoice document based on the provided details.
    This function is adapted from the main() function in invoice_generator.py.
    """
    logger.info(f"Starting invoice generation for invoice #{details['invoice_number']} at {output_path}")
    logger.debug(f"Invoice details: {details}")
    
    # Create and configure document
    doc = Document()
    logger.debug("Created new Document object")
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = details['font_name']
    font.size = Pt(11)
    logger.debug(f"Set default font to {details['font_name']} with size 11pt")

    # Header with icon placeholder and invoice/client details
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(details['column_widths'][0])
    table.columns[1].width = Inches(details['column_widths'][1])
    logger.debug(f"Created header table with column widths {details['column_widths']}")

    # Left cell - Company Icon Placeholder
    cell_left = table.cell(0, 0)
    cell_left.text = '' # Clear placeholder text
    
    # Add image - handle both file path and base64 data
    try:
        # First check if we have icon_data (base64)
        if 'icon_data' in details and details['icon_data']:
            logger.info(f"Using provided base64 icon data for {details['icon_name']}")
            import base64
            import tempfile
            import re
            
            # Extract the actual base64 content if it has a data URL prefix
            base64_data = details['icon_data']
            if base64_data.startswith('data:'):
                # Extract the base64 part after the comma
                match = re.match(r'data:image/[^;]+;base64,(.+)', base64_data)
                if match:
                    base64_data = match.group(1)
                else:
                    logger.warning(f"Couldn't parse base64 data URL format")
            
            # Create a temporary file with the icon data
            with tempfile.NamedTemporaryFile(delete=False, suffix='.' + details['icon_name'].split('.')[-1]) as temp_icon:
                temp_icon.write(base64.b64decode(base64_data))
                temp_icon_path = temp_icon.name
                
            logger.info(f"Created temporary icon file at: {temp_icon_path}")
            cell_left.paragraphs[0].add_run().add_picture(temp_icon_path, width=Inches(2.0))
            logger.info(f"Successfully added company icon from base64 data")
            
            # Clean up the temporary file
            try:
                os.unlink(temp_icon_path)
                logger.debug(f"Removed temporary icon file: {temp_icon_path}")
            except Exception as e:
                logger.warning(f"Could not remove temporary icon file: {e}")
        
        # If no icon_data, try to find the file by name
        else:
            # First try with the exact name provided
            icon_path = details['icon_name']
            logger.info(f"Attempting to load icon from: {icon_path}")
            
            if not os.path.exists(icon_path):
                # If not found, check in the current directory
                current_dir = os.path.dirname(os.path.abspath(__file__))
                icon_path = os.path.join(current_dir, details['icon_name'])
                logger.info(f"Icon not found at direct path, trying current directory: {icon_path}")
                
                if not os.path.exists(icon_path):
                    # Also check in the invoices directory
                    icon_path = os.path.join(invoices_dir, details['icon_name'])
                    logger.info(f"Icon not found in current directory, trying invoices directory: {icon_path}")
                    
                    if not os.path.exists(icon_path):
                        raise FileNotFoundError(f"Icon file not found in any location: {details['icon_name']}")
            
            cell_left.paragraphs[0].add_run().add_picture(icon_path, width=Inches(2.0))
            logger.info(f"Successfully added company icon from: {icon_path}")
    except Exception as e:
        logger.warning(f"Could not add company icon: {str(e)}")
        logger.warning(f"Attempted to find icon at: {details['icon_name']}")
        icon_missing_text = f"{details['company_name']} (Icon not found)"
        cell_left.text = icon_missing_text
        logger.debug(f"Using text placeholder for icon: {icon_missing_text}")
        logger.warning("Make sure the icon file is available in the backend directory or provide a full path")

    # Right cell - Invoice Number and Client Info
    cell_right = table.cell(0, 1)

    # Add Company Name Heading
    p_company_name = cell_right.paragraphs[0] # Use the first paragraph for the company name
    p_company_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_company_name = p_company_name.add_run(details['company_name'])
    run_company_name.bold = True
    run_company_name.font.size = Pt(16)
    logger.debug(f"Added company name: {details['company_name']}")

    # Get date from provided details or use today's date
    from datetime import date, datetime
    try:
        if 'invoice_date' in details and details['invoice_date']:
            # Try to parse the date from input
            parsed_date = datetime.strptime(details['invoice_date'], '%d.%m.%y')
            today_date_str = parsed_date.strftime('%d.%m.%y')
            logger.info(f"Using provided date: {today_date_str}")
        else:
            today_date_str = date.today().strftime('%d.%m.%y')
            logger.info(f"Using today's date: {today_date_str}")
    except (ValueError, TypeError) as e:
        # If date parsing fails, use today's date
        today_date_str = date.today().strftime('%d.%m.%y')
        logger.warning(f"Failed to parse date ({e}), using today's date: {today_date_str}")

    # Add Invoice details in a new paragraph below the company name
    p_invoice_details = cell_right.add_paragraph() # Add a new paragraph for the rest
    p_invoice_details.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Clear existing content before adding runs
    p_invoice_details.text = '' 

    # Add runs piece by piece to apply highlighting and formatting
    run = p_invoice_details.add_run("Invoice #: ")
    run.bold = True
    run = p_invoice_details.add_run(str(details['invoice_number']))
    run.bold = True

    run = p_invoice_details.add_run(f"\nDate: {today_date_str}\n\n")
    run.bold = True

    run = p_invoice_details.add_run(details['client_name'])
    run.bold = True

    run = p_invoice_details.add_run("\n") # Keep the newline separate
    run = p_invoice_details.add_run(details['client_address'])
    run.bold = True
    logger.debug(f"Added invoice header with number {details['invoice_number']} and client {details['client_name']}")

    doc.add_paragraph()  # space

    # Table for services
    doc.add_paragraph("Invoice Details", style='Heading 2')

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    # Set custom column widths: Date, Description, Total
    table.columns[0].width = Inches(1.0)
    table.columns[1].width = Inches(4.0)
    table.columns[2].width = Inches(1.0)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Description of Service'
    hdr_cells[2].text = 'Total'
    hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logger.debug("Created services table with headers")

    # Add background color to header cells
    light_green_color = "A9D08E"
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), light_green_color)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    logger.debug(f"Applied background color {light_green_color} to header cells")

    # Process services and calculate costs
    import re
    subtotal = 0
    logger.info(f"Processing {len(details['services'])} services and calculating costs")
    
    for service in details['services']:
        row_cells = table.add_row().cells
        
        # Extract date from service description using regex
        date_match = re.search(r'(\d{1,2}\.\d{1,2}\.\d{2,4})', service)
        date_str = date_match.group(1) if date_match else ""
        
        # Extract hours from service description
        hours_match = re.search(r'\((\d+\.?\d*)\s*hours?\)', service)
        hours = float(hours_match.group(1)) if hours_match else 0
        
        # Create a clean description without the date
        description = service
        if date_match:
            description = re.sub(r'\s*\d{1,2}\.\d{1,2}\.\d{2,4}\s*', ' ', description).strip()
        
        # Fill cells
        date_cell = row_cells[0]
        date_cell.text = date_str
        date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        desc_cell = row_cells[1]
        desc_cell.text = description
        
        # Calculate cost
        cost = hours * details['hourly_rate']
        row_cells[2].text = f'£{cost:.2f}'
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        subtotal += cost
        logger.info(f"Service: {service} - Date: {date_str} - Hours: {hours} - Cost: £{cost:.2f}")

    logger.info(f"Subtotal calculated: £{subtotal:.2f}")

    # Totals
    row_subtotal = table.add_row().cells
    row_subtotal[0].text = ''
    row_subtotal[1].text = 'Subtotal'
    row_subtotal[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row_subtotal[2].text = f'£{subtotal:.2f}'
    row_subtotal[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logger.debug("Added subtotal row")

    vat_amount = subtotal * (details['vat_rate'] / 100)
    logger.info(f"VAT amount calculated ({details['vat_rate']}%): £{vat_amount:.2f}")

    row_vat = table.add_row().cells
    row_vat[0].text = ''
    row_vat[1].text = f'VAT ({details["vat_rate"]}%)'
    row_vat[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row_vat[2].text = f'£{vat_amount:.2f}'
    row_vat[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logger.debug(f"Added VAT row with rate {details['vat_rate']}%")

    total = subtotal + vat_amount
    logger.info(f"Total amount due: £{total:.2f}")

    row_total = table.add_row().cells
    # Make both cells of Total Amount Due bold
    row_total[0].text = ''
    
    total_due_cell = row_total[1]
    total_due_cell.text = ''  # Clear existing content
    total_due_run = total_due_cell.paragraphs[0].add_run('Total Amount Due')
    total_due_run.bold = True
    total_due_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Make the amount bold too
    total_amount_cell = row_total[2]
    total_amount_cell.text = ''  # Clear existing content
    total_amount_run = total_amount_cell.paragraphs[0].add_run(f'£{total:.2f}')
    total_amount_run.bold = True
    total_amount_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logger.debug("Added total amount due row with bold formatting")

    doc.add_paragraph()  # space

    # Footer details
    footer_text = (
        f"Payment terms of within {details['payment_terms_days']} days.\n"
        "Please make payment by direct transfer to:\n"
        f"Bank Address: {details['bank_address']}\n"
        f"Account name: {details['company_name']}\n"
        f"Account Number: {details['account_number']}\n"
        f"Sort Code: {details['sort_code']}\n"
        "\n"
        "Thanks for your business!\n"
        "\n"
        f"{details['company_name']}, Registered in the UK. Company Number: {details['company_number']}\n"
        f"Registered office: {details['registered_address']}\n"
        f"Registered for VAT in the UK. Registration number: {details['vat_number']}\n"
        f"Email: {details['email']} | Contact: {details['contact_number']}"
    )

    footer_paragraph = doc.add_paragraph()
    run = footer_paragraph.add_run(footer_text)
    run.font.size = Pt(8)
    logger.debug("Added footer with payment and company details")

    # Add PAID stamp as a watermark in the footer, bottom left, if paid
    if details.get('paid'):
        try:
            import os
            current_dir = os.path.dirname(os.path.abspath(__file__))
            stamp_path = os.path.join(current_dir, 'paid_stamp.png')
            if os.path.exists(stamp_path):
                section = doc.sections[0]
                footer = section.footer
                # Remove existing paragraphs in footer
                for p in footer.paragraphs:
                    p.clear()
                # Add a new paragraph for the watermark
                paragraph = footer.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(stamp_path, width=Inches(1.5))
                # Optionally, set paragraph spacing to 0
                paragraph.paragraph_format.space_after = 0
                paragraph.paragraph_format.space_before = 0
                # The image will be in the footer, bottom left, and appear behind content
            else:
                logger.warning(f"PAID stamp image not found at {stamp_path}")
        except Exception as e:
            logger.error(f"Failed to add PAID watermark to footer: {e}")

    # Save the document
    logger.info(f"Saving invoice document to: {output_path}")
    doc.save(output_path)
    logger.info(f"Invoice document saved successfully")

    # Return PDF path if requested
    if generate_pdf:
        logger.info(f"Converting document to PDF format")
        pdf_path = convert_to_pdf(output_path, logger)
        logger.info(f"PDF conversion complete, output at: {pdf_path}")
        return pdf_path
    
    return output_path

@app.post("/generate-invoice", 
    response_class=FileResponse,
    summary="Generate invoice document",
    description="""
Generate an invoice in DOCX or PDF format based on the provided details.

The endpoint accepts a JSON object containing all the necessary invoice information 
and returns a downloadable document file.

The invoice table includes three columns:
- Date: Extracted from the service description (DD.MM.YY format)
- Description of Service: The service provided
- Total: The cost calculated based on hours × hourly rate

Invoice calculation follows these rules:
- Each service cost = hours × hourly rate
- Subtotal = sum of all service costs 
- VAT amount = subtotal × (VAT rate ÷ 100)
- Total amount due = subtotal + VAT amount

Hours are automatically extracted from service descriptions, which must include 
the time in parentheses, e.g., "Consulting Service 21.04.25 (2 hours)".
The date should be in DD.MM.YY format within the service description.
    """,
    response_description="The generated invoice file for download",
    status_code=status.HTTP_200_OK,
    tags=["Invoices"],
    responses={
        200: {
            "content": {
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {},
                "application/pdf": {},
            },
            "description": "Return the generated invoice file",
        },
        400: {
            "description": "Invalid input data or format parameter",
            "content": {
                "application/json": {
                    "example": {"detail": "Invalid format. Must be 'docx' or 'pdf'"}
                }
            }
        },
        500: {
            "description": "Internal server error",
            "content": {
                "application/json": {
                    "example": {"detail": "Failed to generate invoice: Error message details"}
                }
            }
        }
    }
)
async def generate_invoice(
    invoice_details: InvoiceDetails,
    format: OutputFormat = Query(OutputFormat.DOCX, description="Output format for the invoice")
):
    """
    Generate an invoice based on the provided details.
    Returns the invoice file for download.
    
    - **invoice_details**: JSON object containing all invoice information
    - **format**: Output format, either 'docx' or 'pdf' (default: 'docx')
    """
    try:
        logger.info(f"Received request to generate invoice #{invoice_details.invoice_number} in {format} format")
        logger.debug(f"Invoice details: {invoice_details.dict()}")
        
        generate_pdf = format.lower() == "pdf"
        
        # Create a temporary YAML file to store the invoice details
        with tempfile.NamedTemporaryFile(suffix='.yaml', delete=False, mode='w') as temp_yaml:
            # Convert the invoice details to a dictionary and include service_date and service_description
            invoice_dict = invoice_details.dict()
            if invoice_details.service_date:
                invoice_dict['service_date'] = invoice_details.service_date
            if invoice_details.service_description:
                invoice_dict['service_description'] = invoice_details.service_description
            yaml.dump(invoice_dict, temp_yaml)
            temp_yaml_path = Path(temp_yaml.name)
            logger.debug(f"Created temporary YAML file at: {temp_yaml_path}")
        
        try:
            # Generate the output path for the invoice
            invoice_path = get_output_path(temp_yaml_path)
            logger.info(f"Output path for invoice: {invoice_path}")
            
            # Generate the invoice
            result_path = generate_invoice_document(
                invoice_dict, 
                invoice_path, 
                generate_pdf
            )
            
            # If PDF was requested, get the PDF path
            if generate_pdf:
                result_path = get_pdf_path(invoice_path)
            
            logger.info(f"Returning invoice file: {result_path}")
            
            # Return the file as a response
            return FileResponse(
                path=result_path,
                filename=f"invoice_{invoice_details.invoice_number}.{format.lower()}",
                media_type=f"application/{'pdf' if format.lower() == 'pdf' else 'vnd.openxmlformats-officedocument.wordprocessingml.document'}"
            )
        
        finally:
            # Clean up the temporary YAML file
            if temp_yaml_path.exists():
                os.unlink(temp_yaml_path)
                logger.debug(f"Cleaned up temporary YAML file: {temp_yaml_path}")
                
    except Exception as e:
        error_msg = f"Error generating invoice: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate invoice: {str(e)}")

@app.get("/", 
    summary="API root endpoint",
    description="API root endpoint with basic information and usage instructions",
    tags=["System"],
)
async def root():
    """API root endpoint with usage information."""
    logger.debug("Received request to root endpoint")
    return {
        "message": "Invoice Generator API",
        "usage": "POST /generate-invoice with required invoice details",
        "documentation": "/docs for Swagger UI documentation",
        "version": "1.0.0",
        "verbose_logging": VERBOSE
    }

# Add middleware to log all requests
@app.middleware("http")
async def log_requests(request, call_next):
    logger.debug(f"Request: {request.method} {request.url}")
    response = await call_next(request)
    logger.debug(f"Response status: {response.status_code}")
    return response


# Custom OpenAPI schema generation
def custom_openapi():
    if app.openapi_schema:
        return app.openapi_schema
    
    openapi_schema = get_openapi(
        title="Invoice Generator API",
        version="1.0.0",
        description="""
This API allows you to generate professional invoices in DOCX and PDF formats.

## Features

* Generate invoices with customizable details
* Support for both DOCX and PDF output formats
* Automatic calculation of subtotals, VAT, and total amounts
* Professional formatting with company branding

## Format Requirements

- Service entries must include hours in parentheses (e.g., "AI Consultancy (2 hours)")
- Invoice date must be in DD.MM.YY format
- All monetary values are in GBP (£)

## Calculation Logic

1. Each service cost = hours × hourly rate
2. Subtotal = sum of all service costs
3. VAT amount = subtotal × (VAT rate ÷ 100)
4. Total amount due = subtotal + VAT amount""",
        routes=app.routes,
    )
    
    # Add example values
    app.openapi_schema = openapi_schema
    return app.openapi_schema

app.openapi = custom_openapi

# API version info endpoint - Useful for checking if API is running and checking CORS
@app.get("/version",
    summary="API version information",
    description="Returns the current API version and configuration information",
    tags=["System"]
)
async def version():
    """Return API version and configuration information."""
    logger.debug("Request received for API version information")
    return {
        "api_name": "Invoice Generator API",
        "version": "1.0.0",
        "cors_enabled": True,
        "cors_origins": CORS_ORIGINS,
        "verbose_logging": VERBOSE
    }

# Example of how to use the API from a web page
@app.get("/example-client",
    summary="Example web client",
    description="Returns a simple HTML page with JavaScript code to test the API",
    tags=["Documentation"],
    response_class=FileResponse
)
async def example_client():
    """Serve a simple HTML page with JavaScript to test the API."""
    logger.debug("Serving example client HTML page")
    
    # Create a temporary HTML file with the example code
    with tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w') as temp_html:
        temp_html.write('''
        <!DOCTYPE html>
        <html>
        <head>
            <title>Invoice Generator API Test Client</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                pre { background-color: #f5f5f5; padding: 10px; border-radius: 5px; overflow-x: auto; }
                button { padding: 8px 15px; background-color: #4CAF50; color: white; border: none; 
                         border-radius: 4px; cursor: pointer; margin: 10px 0; }
                button:hover { background-color: #45a049; }
                #result { margin-top: 20px; }
            </style>
        </head>
        <body>
            <h1>Invoice Generator API Test Client</h1>
            <p>This page demonstrates how to call the Invoice Generator API from a web application.</p>
            
            <h2>Test API Connection</h2>
            <button onclick="testConnection()">Test Connection</button>
            
            <h2>Generate Invoice Example</h2>
            <pre>
const invoiceData = {
    "client_name": "Web Client Test",
    "client_address": "123 Browser St.\\nWeb City\\nW1 1AA\\nU.K.",
    "services": ["API Testing 01.05.25 (2 hours)", "Documentation 01.05.25 (1 hour)"],
    "payment_terms_days": 30,
    "invoice_number": 1010,
    "invoice_date": "01.05.25",
    "company_name": "Test Consulting Ltd",
    "hourly_rate": 200,
    "vat_rate": 20,
    "account_number": "12345678",
    "sort_code": "12-34-56",
    "bank_address": "123 Bank St, London, UK",
    "company_number": "12345678",
    "vat_number": "GB123456789",
    "registered_address": "123 Business St, London, UK",
    "email": "contact@testconsulting.com",
    "contact_number": "07700 900123",
    "column_widths": [2.5, 3.5],
    "font_name": "DejaVu Sans",
    "icon_name": "DioramaConsultingIcon.png"
};
            </pre>
            
            <button onclick="generateInvoice()">Generate Invoice (DOCX)</button>
            <button onclick="generatePdfInvoice()">Generate Invoice (PDF)</button>
            
            <div id="result"></div>
            
            <script>
                // Function to test API connection
                async function testConnection() {
                    const resultDiv = document.getElementById('result');
                    resultDiv.innerHTML = 'Testing connection...';
                    
                    try {
                        const response = await fetch('http://localhost:8083/version');
                        if (response.ok) {
                            const data = await response.json();
                            resultDiv.innerHTML = `<p>✅ Connection successful!</p>
                                <pre>${JSON.stringify(data, null, 2)}</pre>`;
                        } else {
                            resultDiv.innerHTML = `<p>❌ Error: ${response.status} ${response.statusText}</p>`;
                        }
                    } catch (error) {
                        resultDiv.innerHTML = `<p>❌ Connection failed: ${error.message}</p>
                            <p>This may be due to CORS issues if the API server does not allow requests from this origin.</p>`;
                    }
                }
                
                // Function to generate invoice
                async function generateInvoice(format = 'docx') {
                    const resultDiv = document.getElementById('result');
                    resultDiv.innerHTML = `Generating ${format.toUpperCase()} invoice...`;
                    
                    try {
                        const invoiceData = {
                            "client_name": "Web Client Test",
                            "client_address": "123 Browser St.\\nWeb City\\nW1 1AA\\nU.K.",
                            "services": ["API Testing 01.05.25 (2 hours)", "Documentation 01.05.25 (1 hour)"],
                            "payment_terms_days": 30,
                            "invoice_number": 1010,
                            "invoice_date": "01.05.25",
                            "company_name": "Test Consulting Ltd",
                            "hourly_rate": 200,
                            "vat_rate": 20,
                            "account_number": "12345678",
                            "sort_code": "12-34-56",
                            "bank_address": "123 Bank St, London, UK",
                            "company_number": "12345678",
                            "vat_number": "GB123456789",
                            "registered_address": "123 Business St, London, UK",
                            "email": "contact@testconsulting.com",
                            "contact_number": "07700 900123",
                            "column_widths": [2.5, 3.5],
                            "font_name": "DejaVu Sans",
                            "icon_name": "DioramaConsultingIcon.png"
                        };
                        
                        // Direct file download approach (binary data)
                        const response = await fetch(`http://localhost:8083/generate-invoice?format=${format}`, {
                            method: 'POST',
                            headers: {
                                'Content-Type': 'application/json'
                            },
                            body: JSON.stringify(invoiceData)
                        });
                        
                        if (!response.ok) {
                            const errorText = await response.text();
                            throw new Error(`API error: ${response.status} ${response.statusText}\\n${errorText}`);
                        }
                        
                        // Get the blob from the response
                        const blob = await response.blob();
                        
                        // Create a download link and click it
                        const url = window.URL.createObjectURL(blob);
                        const a = document.createElement('a');
                        a.href = url;
                        a.download = `invoice_1010.${format}`;
                        document.body.appendChild(a);
                        a.click();
                        window.URL.revokeObjectURL(url);
                        a.remove();
                        
                        resultDiv.innerHTML = `<p>✅ ${format.toUpperCase()} invoice generated successfully! Check your downloads.</p>`;
                    } catch (error) {
                        resultDiv.innerHTML = `<p>❌ Error generating invoice: ${error.message}</p>`;
                    }
                }
                
                // Function to generate PDF invoice
                function generatePdfInvoice() {
                    generateInvoice('pdf');
                }
            </script>
        </body>
        </html>
        ''')
        temp_html_path = Path(temp_html.name)
    
    return FileResponse(
        path=temp_html_path,
        filename="invoice_api_example.html",
        media_type="text/html"
    )

if __name__ == "__main__":
    import uvicorn # type: ignore
    
    logger.info(f"Starting uvicorn server with VERBOSE={VERBOSE}")
    uvicorn.run(app, host="0.0.0.0", port=8083) 