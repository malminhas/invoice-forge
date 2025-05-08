"""
Invoice Generator Script

Usage:
    invoice_generator.py [-v] [-V] [--pdf] [--pdf-backend=<backend>] <yaml_file>
    invoice_generator.py -h

Options:
    -h --help       Show this help message and exit
    -v --verbose    Enable verbose logging output
    -V --version    Show version information
    --pdf           Generate PDF output in addition to DOCX
    --pdf-backend=<backend>  PDF backend: 'libreoffice' (default) or 'docx2pdf'
    <yaml_file>     Name of YAML configuration file (will be loaded from invoices directory)

This script generates an invoice document using data from a YAML file.

Installation Instructions:
1. Ensure you have Python installed on your system.
2. Install the required libraries using pip:
   - python-docx: `pip install python-docx`
   - PyYAML: `pip install PyYAML`
   - docopt: `pip install docopt`
3. For PDF conversion (if using --pdf flag):
   - Install LibreOffice: `brew install libreoffice`
   - Install unoconv: `brew install unoconv`

Instructions:
- Place the 'DioramaConsultingIcon.png' image in the same directory as this script.
- Place your YAML configuration files in the 'invoices' directory.
- Run the script: python invoice_generator.py client1.yaml
- Add --pdf flag to generate PDF output: python invoice_generator.py --pdf client1.yaml
- Generated invoices will be saved in the 'invoices' directory
- Use -v for detailed logging output

Invoice Calculation Logic:
1. Service Cost Calculation:
   - Each service entry should include hours in parentheses, e.g., "AI Consultancy (2 hours)"
   - The cost per service is calculated as: hours × hourly_rate
   - Example: 2 hours at £100/hour = £200

2. Total Amount Calculation:
   - Subtotal: Sum of all individual service costs
   - VAT Amount: Calculated as (subtotal × VAT_rate)
   - Total Amount Due: subtotal + VAT amount
   
   Example calculation:
   - Services: 2 hours at £100/hour = £200 subtotal
   - VAT at 20%: £200 × 0.20 = £40
   - Total Amount Due: £200 + £40 = £240

Note: All monetary values are formatted to 2 decimal places in the output.
"""

import logging
from docx import Document # type: ignore
from docx.shared import Inches, Pt # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH # type: ignore
from docx.oxml.ns import qn # type: ignore
from docx.oxml import OxmlElement # type: ignore
import os
import subprocess
from datetime import date, datetime
from docx.enum.text import WD_COLOR_INDEX # type: ignore
import yaml # type: ignore
import re
import sys
from pathlib import Path

INVOICES_DIR = 'invoices'

__version__ = '1.0.0'

def setup_logging(verbose):
    """Configure logging based on verbose flag."""
    if verbose:
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
    else:
        logging.basicConfig(level=logging.WARNING)

def ensure_invoices_directory():
    """Ensure the invoices directory exists."""
    invoices_dir = Path(INVOICES_DIR)
    invoices_dir.mkdir(exist_ok=True)
    return invoices_dir

def get_yaml_path(yaml_filename):
    """Get the full path for a YAML file in the invoices directory."""
    invoices_dir = ensure_invoices_directory()
    yaml_path = invoices_dir / yaml_filename
    if not yaml_path.exists():
        raise FileNotFoundError(f"YAML file not found in invoices directory: {yaml_path}")
    return yaml_path

def validate_yaml_path(yaml_path):
    """Validate that the YAML file exists."""
    if not yaml_path.exists():
        raise FileNotFoundError(f"YAML file not found: {yaml_path}")
    return yaml_path

def get_output_path(yaml_path):
    """Generate output path for the invoice based on input YAML filename."""
    yaml_stem = yaml_path.stem  # Get filename without extension
    return yaml_path.parent / f"{yaml_stem}_invoice.docx"

def get_pdf_path(docx_path):
    """Generate PDF path from DOCX path."""
    return docx_path.with_suffix('.pdf')

def convert_to_pdf(docx_path: Path, logger, backend: str = 'libreoffice') -> Path:
    """
    Convert a DOCX file to PDF format using the selected backend.
    Falls back to 'docx2pdf' if 'libreoffice' fails.
    """
    def try_docx2pdf_fallback() -> Path:
        try:
            from docx2pdf import convert
            pdf_path = get_pdf_path(docx_path)
            logger.info("Falling back to docx2pdf for PDF conversion...")
            convert(str(docx_path), str(pdf_path))

            if pdf_path.exists():
                logger.info(f"PDF conversion successful (docx2pdf): {pdf_path}")
                return pdf_path
            else:
                raise RuntimeError("docx2pdf fallback failed: Output file was not created.")
        except ImportError:
            raise RuntimeError("Fallback failed: docx2pdf is not installed. Run 'pip install docx2pdf'")
        except Exception as e:
            raise RuntimeError(f"Fallback to docx2pdf failed: {str(e)}")

    if backend == 'libreoffice':
        try:
            subprocess.run(['which', 'unoconv'], check=True, capture_output=True)

            pdf_path = get_pdf_path(docx_path)
            logger.info(f"Converting DOCX to PDF using LibreOffice/unoconv: {pdf_path}")

            result = subprocess.run(
                [
                    'unoconv',
                    '-f', 'pdf',
                    str(docx_path)
                ],
                capture_output=True,
                text=True
            )

            if result.returncode != 0:
                logger.error(f"unoconv failed with exit code {result.returncode}")
                logger.error(f"unoconv stdout:\n{result.stdout}")
                logger.error(f"unoconv stderr:\n{result.stderr}")
                raise RuntimeError(f"unoconv failed with exit code {result.returncode}")

            if not pdf_path.exists():
                logger.error(f"unoconv ran but did not create output.\nstdout:\n{result.stdout}\nstderr:\n{result.stderr}")
                raise RuntimeError(f"unoconv ran but did not create output.\nstdout:\n{result.stdout}\nstderr:\n{result.stderr}")

            logger.info("PDF conversion completed successfully (LibreOffice)")
            return pdf_path

        except Exception as e:
            logger.warning(f"LibreOffice backend failed: {str(e)}")
            return try_docx2pdf_fallback()

    elif backend == 'docx2pdf':
        return try_docx2pdf_fallback()

    else:
        error_msg = f"Unknown PDF backend: {backend}. Supported: 'libreoffice', 'docx2pdf'"
        logger.error(error_msg)
        raise RuntimeError(error_msg)

'''
def convert_to_pdf(docx_path: Path, logger, backend: str = 'libreoffice') -> Path:
    """
    Convert a DOCX file to PDF format using the selected backend.
    Args:
        docx_path (Path): Path to the DOCX file
        logger: Logger instance for logging messages
        backend (str): 'libreoffice' or 'docx2pdf'
    Returns:
        Path: Path to the generated PDF file
    Raises:
        RuntimeError: If PDF conversion fails
    """
    if backend == 'libreoffice':
        try:
            # Check if unoconv is installed
            subprocess.run(['which', 'unoconv'], check=True, capture_output=True)
            pdf_path = get_pdf_path(docx_path)
            logger.info(f"Converting DOCX to PDF using LibreOffice/unoconv: {pdf_path}")
            result = subprocess.run(['unoconv', '-f', 'pdf', str(docx_path)], check=True, capture_output=True, text=True)
            if not pdf_path.exists():
                raise RuntimeError(f"PDF conversion failed. unoconv output: {result.stderr}")
            logger.info("PDF conversion completed successfully (LibreOffice)")
            return pdf_path
        except subprocess.CalledProcessError:
            error_msg = "unoconv is not installed or failed. Please install it using 'brew install unoconv' or check your LibreOffice installation."
            logger.error(error_msg)
            raise RuntimeError(error_msg)
        except Exception as e:
            logger.error(f"Failed to convert to PDF (LibreOffice): {str(e)}")
            raise
    elif backend == 'docx2pdf':
        try:
            from docx2pdf import convert
            pdf_path = get_pdf_path(docx_path)
            logger.info(f"Converting {docx_path} to PDF using docx2pdf...")
            convert(str(docx_path), str(pdf_path))
            if pdf_path.exists():
                logger.info(f"PDF conversion successful: {pdf_path}")
                return pdf_path
            else:
                error_msg = "PDF conversion failed: Output file was not created"
                logger.error(error_msg)
                raise RuntimeError(error_msg)
        except ImportError:
            error_msg = "docx2pdf is not installed. Please install it using 'pip install docx2pdf'"
            logger.error(error_msg)
            raise RuntimeError(error_msg)
        except Exception as e:
            error_msg = f"PDF conversion failed (docx2pdf): {str(e)}"
            logger.error(error_msg)
            raise RuntimeError(error_msg)
    else:
        error_msg = f"Unknown PDF backend: {backend}. Supported: 'libreoffice', 'docx2pdf'"
        logger.error(error_msg)
        raise RuntimeError(error_msg)
'''

def load_details(file_path):
    """Load and validate configuration from YAML file."""
    logger = logging.getLogger(__name__)
    logger.info(f"Loading configuration from {file_path}")
    try:
        with open(file_path, 'r') as file:
            details = yaml.safe_load(file)
            
        # Log all the values from YAML only if in verbose mode
        if logger.isEnabledFor(logging.INFO):
            logger.info("Successfully loaded YAML configuration:")
            logger.info(f"Company Name: {details.get('company_name')}")
            logger.info(f"Invoice Number: {details.get('invoice_number')}")
            if 'invoice_date' in details:
                logger.info(f"Invoice Date: {details.get('invoice_date')}")
            logger.info(f"Client Name: {details.get('client_name')}")
            logger.info(f"Client Address: {details.get('client_address')}")
            logger.info(f"Services: {details.get('services')}")
            logger.info(f"Hourly Rate: £{details.get('hourly_rate')}")
            logger.info(f"VAT Rate: {details.get('vat_rate')}%")
            logger.info(f"Payment Terms: {details.get('payment_terms_days')} days")
            logger.info(f"Bank Details:")
            logger.info(f"  - Account Number: {details.get('account_number')}")
            logger.info(f"  - Sort Code: {details.get('sort_code')}")
            logger.info(f"  - Bank Address: {details.get('bank_address')}")
            logger.info(f"Company Details:")
            logger.info(f"  - Company Number: {details.get('company_number')}")
            logger.info(f"  - VAT Number: {details.get('vat_number')}")
            logger.info(f"  - Registered Address: {details.get('registered_address')}")
            logger.info(f"Contact Details:")
            logger.info(f"  - Email: {details.get('email')}")
            logger.info(f"  - Phone: {details.get('contact_number')}")
            logger.info(f"Document Settings:")
            logger.info(f"  - Font Name: {details.get('font_name')}")
            logger.info(f"  - Icon Name: {details.get('icon_name')}")
            logger.info(f"  - Column Widths: {details.get('column_widths')}")
        
        return details
    except FileNotFoundError:
        logger.error(f"Configuration file {file_path} not found")
        raise
    except yaml.YAMLError as e:
        logger.error(f"Error parsing YAML file: {e}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error loading configuration: {e}")
        raise

def main():
    from docopt import docopt # type: ignore

    """Main entry point for the invoice generator."""
    arguments = docopt(__doc__, version=f'Invoice Generator v{__version__}')
    
    # Setup logging based on verbose flag
    setup_logging(arguments['--verbose'])
    
    logger = logging.getLogger(__name__)
    
    try:
        # Ensure invoices directory exists and get YAML path
        yaml_path = get_yaml_path(arguments['<yaml_file>'])
        
        # Get the output path based on input filename
        output_path = get_output_path(yaml_path)
        
        # Load details from specified YAML file
        details = load_details(yaml_path)
        logger.info("Starting invoice generation...")
        
        # Create and configure document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = details['font_name']
        font.size = Pt(12)

        # Header with icon placeholder and invoice/client details
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(details['column_widths'][0])
        table.columns[1].width = Inches(details['column_widths'][1])

        # Left cell - Company Icon Placeholder
        cell_left = table.cell(0, 0)
        cell_left.text = '' # Clear placeholder text
        # Add image - ensure the icon is in the same directory or provide full path
        cell_left.paragraphs[0].add_run().add_picture(details['icon_name'], width=Inches(2.0))

        # Right cell - Invoice Number and Client Info (placeholders)
        cell_right = table.cell(0, 1)

        # Add Company Name Heading
        p_company_name = cell_right.paragraphs[0] # Use the first paragraph for the company name
        p_company_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_company_name = p_company_name.add_run(details['company_name'])
        run_company_name.bold = True
        run_company_name.font.size = Pt(16)

        # Get date - either from YAML or today's date
        try:
            if 'invoice_date' in details:
                # Try to parse the date from YAML
                parsed_date = datetime.strptime(details['invoice_date'], '%d.%m.%y')
                today_date_str = parsed_date.strftime('%d.%m.%y')
                logger.info(f"Using date from YAML file: {today_date_str}")
            else:
                today_date_str = date.today().strftime('%d.%m.%y')
                logger.info(f"Using today's date: {today_date_str}")
        except (ValueError, TypeError):
            # If date parsing fails, use today's date
            today_date_str = date.today().strftime('%d.%m.%y')
            logger.warning(f"Failed to parse date from YAML, using today's date: {today_date_str}")

        # Add Invoice details in a new paragraph below the company name
        p_invoice_details = cell_right.add_paragraph() # Add a new paragraph for the rest
        p_invoice_details.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Clear existing content before adding runs
        p_invoice_details.text = '' 

        # Add runs piece by piece to apply highlighting and formatting
        run = p_invoice_details.add_run("Invoice #: ")
        run.bold = True
        run = p_invoice_details.add_run(str(details['invoice_number']))
        run.bold = True

        run = p_invoice_details.add_run(f"\nDate: {today_date_str}\n\n")
        run.bold = True

        run = p_invoice_details.add_run(details['client_name'])
        run.bold = True

        run = p_invoice_details.add_run("\n") # Keep the newline separate
        run = p_invoice_details.add_run(details['client_address'])
        run.bold = True

        doc.add_paragraph()  # space

        # Table for services
        doc.add_paragraph("Invoice Details", style='Heading 2')

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Description of Service'
        hdr_cells[1].text = 'Total'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add background color to header cells
        light_green_color = "A9D08E"
        for cell in hdr_cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), light_green_color)
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)

        # Placeholder rows for services
        subtotal = 0
        if logger.isEnabledFor(logging.INFO):
            logger.info("Processing services and calculating costs:")
        for service in details['services']:
            row_cells = table.add_row().cells
            desc_cell = row_cells[0]
            desc_cell.text = '' # Clear cell
            run = desc_cell.paragraphs[0].add_run(service)
            
            # Extract hours from service description (e.g., "AI Consultancy (1 hour)" -> 1)
            hours_match = re.search(r'\((\d+\.?\d*)\s*hours?\)', service)
            hours = float(hours_match.group(1)) if hours_match else 0
            
            # Calculate cost
            cost = hours * details['hourly_rate']
            row_cells[1].text = f'£{cost:.2f}'
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            subtotal += cost
            if logger.isEnabledFor(logging.INFO):
                logger.info(f"Service: {service} - Hours: {hours} - Cost: £{cost:.2f}")

        if logger.isEnabledFor(logging.INFO):
            logger.info(f"Subtotal calculated: £{subtotal:.2f}")

        # Totals
        row_subtotal = table.add_row().cells
        row_subtotal[0].text = 'Subtotal'
        row_subtotal[1].text = f'£{subtotal:.2f}'
        row_subtotal[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        vat_amount = subtotal * (details['vat_rate'] / 100)
        if logger.isEnabledFor(logging.INFO):
            logger.info(f"VAT amount calculated ({details['vat_rate']}%): £{vat_amount:.2f}")

        row_vat = table.add_row().cells
        row_vat[0].text = f'VAT ({details["vat_rate"]}%)'
        row_vat[1].text = f'£{vat_amount:.2f}'
        row_vat[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        total = subtotal + vat_amount
        if logger.isEnabledFor(logging.INFO):
            logger.info(f"Total amount due: £{total:.2f}")

        row_total = table.add_row().cells
        # Make both cells of Total Amount Due bold
        total_due_cell = row_total[0]
        total_due_cell.text = '' # Clear existing content
        total_due_run = total_due_cell.paragraphs[0].add_run('Total Amount Due')
        total_due_run.bold = True

        # Make the amount bold too
        total_amount_cell = row_total[1]
        total_amount_cell.text = '' # Clear existing content
        total_amount_run = total_amount_cell.paragraphs[0].add_run(f'£{total:.2f}')
        total_amount_run.bold = True
        total_amount_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()  # space

        # Footer details
        footer_text = (
            f"Payment terms of within {details['payment_terms_days']} days.\n"
            "Please make payment by direct transfer to:\n"
            f"Bank Address: {details['bank_address']}\n"
            f"Account name: {details['company_name']}\n"
            f"Account Number: {details['account_number']}\n"
            f"Sort Code: {details['sort_code']}\n"
            "\n"
            "Thanks for your business!\n"
            "\n"
            f"{details['company_name']}, Registered in the UK. Company Number: {details['company_number']}\n"
            f"Registered office: {details['registered_address']}\n"
            f"Registered for VAT in the UK. Registration number: {details['vat_number']}\n"
            f"Email: {details['email']} | Contact: {details['contact_number']}"
        )

        footer_paragraph = doc.add_paragraph()
        run = footer_paragraph.add_run(footer_text)
        run.font.size = Pt(8)

        # Add PAID stamp as a watermark in the footer, bottom left, if paid
        if details.get('paid'):
            try:
                current_dir = os.path.dirname(os.path.abspath(__file__))
                stamp_path = os.path.join(current_dir, 'paid_stamp.png')
                if os.path.exists(stamp_path):
                    section = doc.sections[0]
                    footer = section.footer
                    # Add a new paragraph for the paid stamp image
                    paragraph = footer.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(stamp_path, width=Inches(1.5))
                    paragraph.paragraph_format.space_after = 0
                    paragraph.paragraph_format.space_before = 0
                else:
                    logger.warning(f"PAID stamp image not found at {stamp_path}")
            except Exception as e:
                logger.error(f"Failed to add PAID watermark to footer: {e}")

        # Save the document with the new output path
        if logger.isEnabledFor(logging.INFO):
            logger.info(f"Saving invoice document to: {output_path}")
        doc.save(output_path)

        # Convert to PDF if requested
        if arguments['--pdf']:
            pdf_backend = arguments.get('--pdf-backend') or 'libreoffice'
            pdf_path = convert_to_pdf(output_path, logger, backend=pdf_backend)
            # Open the PDF instead of DOCX if it was generated
            output_path = pdf_path

        # Open the generated document on macOS
        try:
            if logger.isEnabledFor(logging.INFO):
                logger.info(f"Attempting to open the document: {output_path}")
            os.system(f"open '{output_path}'")
            if logger.isEnabledFor(logging.INFO):
                logger.info("Document opened successfully")
        except Exception as e:
            logger.error(f"Could not automatically open the file. Error: {e}")

        if logger.isEnabledFor(logging.INFO):
            logger.info("Invoice generation completed successfully")
        
    except Exception as e:
        logger.error(f"Error generating invoice: {str(e)}")
        sys.exit(1)

if __name__ == '__main__':
    main()

